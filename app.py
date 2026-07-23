import streamlit as st
import pandas as pd
import io
import re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
import requests
import json
# ==========================================
# 🛠️ FUNÇÕES DE FORMATAÇÃO CEMT
# ==========================================
def extrair_valor(row, mapping):
    if mapping["tipo"] == "nenhum":
        return ""
    elif mapping["tipo"] == "fixo":
        return str(mapping["valor"]).strip()
    elif mapping["tipo"] == "coluna" and mapping["valor"] != "--- Nenhuma ---":
        val = row[mapping["valor"]]
        return "" if pd.isna(val) else str(val).strip()
    return ""

def formatar_data_romana(dt_ini, dt_fim):
    if not dt_ini: return ""
    try:
        meses_romanos = ['i', 'ii', 'iii', 'iv', 'v', 'vi', 'vii', 'viii', 'ix', 'x', 'xi', 'xii']
        d1 = pd.to_datetime(dt_ini, dayfirst=True)
        m1_rom = meses_romanos[d1.month - 1]
        
        if dt_fim:
            d2 = pd.to_datetime(dt_fim, dayfirst=True)
            if d1.month == d2.month and d1.year == d2.year:
                if d1.day == d2.day:
                    return f"{d1.day:02d}.{m1_rom}.{d1.year}"
                return f"{d1.day:02d}-{d2.day:02d}.{m1_rom}.{d1.year}"
            else:
                m2_rom = meses_romanos[d2.month - 1]
                return f"{d1.day:02d}.{m1_rom}.{d1.year}-{d2.day:02d}.{m2_rom}.{d2.year}"
        return f"{d1.day:02d}.{m1_rom}.{d1.year}"
    except:
        return str(dt_ini)

def remover_milesimos_segundos(coord):
    if not coord: return ""
    return re.sub(r'(\d+)\.\d+(["”\'’])', r'\1\2', str(coord))

# ==========================================
# 🎨 CONFIGURAÇÃO DA PÁGINA E ESTILO VISUAL
# ==========================================
st.set_page_config(page_title="Gerador Universal", page_icon="🪲", layout="wide")

# ==========================================
# 🎨 CONFIGURAÇÃO DA PÁGINA E ESTILO VISUAL
# ==========================================
st.set_page_config(page_title="Gerador Universal CEMT", page_icon="🪲", layout="wide")

st.markdown("""
    <style>
        /* 1. Título principal (st.title) */
        h1 {
            font-size: 48px !important;
        }

        /* 2. Subtítulos da página (st.subheader) */
        h3 {
            font-size: 32px !important;
        }

        /* 3. Texto comum do corpo (st.write, st.markdown e frases explicativas) */
        div[data-testid="stMarkdownContainer"] p {
            font-size: 24px !important;
        }

        /* 4. Texto dos balões coloridos de aviso (st.success, st.info, st.error) */
        .stAlert p {
            font-size: 20px !important;
        }

        /* 5. Rótulos/Perguntas acima dos campos */
        .stSidebar label, .stRadio label, .stSelectbox label, .stTextInput label, .stCheckbox label {
            font-size: 22px !important;
            font-weight: bold;
        }

        /* 6. 🟢 TEXTO DENTRO DAS CAIXAS DE SELEÇÃO E CAMPOS DE TEXTO */
        .stTextInput input, div[data-baseweb="select"] div {
            font-size: 22px !important;
        }

        /* 7. 🟢 TEXTO DAS OPÇÕES DO MENU SUSPENSO QUANDO ABRIR A CAIXA */
        ul[data-baseweb="menu"] li span, div[data-baseweb="popover"] span {
            font-size: 22px !important;
        }

        /* 8. 🟢 TEXTO DOS BOTÕES DE RÁDIO (Usar coluna / Valor fixo / Não incluir) */
        div[role="radiogroup"] p {
            font-size: 19px !important;
        }

        /* 9. Título das caixas retráteis (st.expander) */
        .stExpander details summary {
            font-size: 20px !important;
        }
        /* 10. 🟢 TEXTO DENTRO DA CAIXA DE SUGESTÕES (st.text_area) */
        .stTextArea textarea {
            font-size: 20px !important;
        }
    </style>
""", unsafe_allow_html=True)

st.title("🪲 Gerador de Etiquetas Entomológicas")

st.subheader("1. Carregue os Dados")
arquivo_upload = st.file_uploader("Arraste ou escolha a planilha Excel (.xlsx)", type=["xlsx"])

if arquivo_upload is not None:
    xls = pd.ExcelFile(arquivo_upload)
    
    col1, col2 = st.columns(2)
    with col1:
        aba_selecionada = st.selectbox("Selecione a aba da planilha:", xls.sheet_names)
    with col2:
        linha_cabecalho = st.number_input("Em qual linha estão os nomes das colunas?", min_value=1, value=1)
    
    df = pd.read_excel(arquivo_upload, sheet_name=aba_selecionada, header=linha_cabecalho - 1)
    colunas_reais = [str(c) for c in df.columns.tolist()]
    lista_colunas = ["--- Nenhuma ---"] + colunas_reais
    
    st.success(f"✅ Encontrados {len(df)} registros totais na aba '{aba_selecionada}'.")
    
    with st.expander("👀 Ver prévia dos dados lidos", expanded=False):
        st.dataframe(df.head(3))

    def mapear_campo(nome_exibicao, chave, tipo_padrao="coluna"):
        st.markdown(f"**{nome_exibicao}**")
        opcoes = ["Usar coluna", "Valor fixo", "Não incluir"]
        if tipo_padrao == "coluna": idx = 0
        elif tipo_padrao == "fixo": idx = 1
        else: idx = 2
        
        modo = st.radio(f"Origem:", opcoes, key=f"modo_{chave}", horizontal=True, index=idx)
        
        if modo == "Usar coluna":
            coluna = st.selectbox("Coluna:", lista_colunas, key=f"col_{chave}")
            st.markdown("---")
            return {"tipo": "coluna", "valor": coluna}
        elif modo == "Valor fixo":
            texto = st.text_input("Valor fixo:", key=f"fixo_{chave}")
            st.markdown("---")
            return {"tipo": "fixo", "valor": texto}
        else:
            st.markdown("---")
            return {"tipo": "nenhum", "valor": ""}

    # ==========================================
    # 🎛️ BARRA LATERAL (MAPEAMENTO E FILTRO)
    # ==========================================
    st.sidebar.header("🗺️ Configurações")

    # 🟢 NOVO: ABA DE FILTRAGEM
    with st.sidebar.expander("🎯 Filtro de Seleção", expanded=True):
        st.write("Gere etiquetas apenas para linhas específicas.")
        ativar_filtro = st.checkbox("Ativar filtro de indivíduos")
        
        coluna_filtro = None
        texto_filtro = ""
        
        if ativar_filtro:
            coluna_filtro = st.selectbox("Filtrar pela coluna:", colunas_reais)
            texto_filtro = st.text_input("Que contenha o texto (separe por vírgulas para mais de um):", placeholder="Ex: referência, imprimir, x")

    with st.sidebar.expander("📍 Localização", expanded=False):
        map_pais = mapear_campo("País", "pais", "fixo")
        map_estado = mapear_campo("Estado / UF", "estado")
        map_mun = mapear_campo("Município", "mun")
        map_loc = mapear_campo("Localidade Específica", "loc")
        map_lat = mapear_campo("Latitude", "lat")
        map_lon = mapear_campo("Longitude", "lon")
        map_alt = mapear_campo("Altitude", "alt", "nenhum") 

    with st.sidebar.expander("📅 Coleta e Armadilha", expanded=False):
        map_dt_ini = mapear_campo("Data Inicial", "dt_ini")
        map_dt_fim = mapear_campo("Data Final", "dt_fim", "nenhum") 
        map_isca = mapear_campo("Armadilha / Isca", "isca", "fixo")
        map_coletor = mapear_campo("Coletor", "coletor", "fixo")
        map_id = mapear_campo("Código da Amostra", "id")

    with st.sidebar.expander("🔬 Taxonomia (Opcional)", expanded=False):
        exibir_titulo_tax = st.checkbox("Exibir identificação no topo?", value=False)
        map_titulo_tax = mapear_campo("Título Taxonômico", "titulo_tax") if exibir_titulo_tax else {"tipo": "nenhum", "valor": ""}

# ==========================================
    # 🚀 GERAÇÃO DO DOCUMENTO WORD
    # ==========================================
    st.subheader("2. Gerar Documento")
    if st.button("🚀 Gerar Etiquetas em Word"):
        
        # 🟢 1. LÓGICA DE FILTRAGEM DOS DADOS (AGORA COM MÚLTIPLOS TERMOS)
        df_filtrado = df.copy()
        if ativar_filtro and texto_filtro.strip():
            # Pega no texto "FLO, PAST, CAF", separa pelas vírgulas e tira os espaços extra
            termos = [re.escape(termo.strip()) for termo in texto_filtro.split(",") if termo.strip()]
            
            # Junta tudo com a barra vertical "OU" (ex: FLO|PAST|CAF)
            padrao_regex = "|".join(termos)
            
            # Filtra o Excel: se a célula contiver QUALQUER UM dos termos, a linha é mantida
            df_filtrado = df_filtrado[df_filtrado[coluna_filtro].astype(str).str.contains(padrao_regex, case=False, na=False, regex=True)]
        
        if len(df_filtrado) == 0:
            st.error(f"❌ Nenhum registo encontrado na coluna '{coluna_filtro}' contendo os termos informados. Verifique a planilha ou o filtro.")
        else:
            with st.spinner(f"Construindo tabela no Word para {len(df_filtrado)} amostras..."):
                
                doc = Document()
                
                # Configurando a página para PAISAGEM e margens de 0.5cm
                for section in doc.sections:
                    section.orientation = WD_ORIENTATION.LANDSCAPE
                    section.page_width = Cm(29.7)
                    section.page_height = Cm(21.0)
                    
                    section.top_margin = Cm(0.5)
                    section.bottom_margin = Cm(0.5)
                    section.left_margin = Cm(0.5)
                    section.right_margin = Cm(0.5)

                # Construindo a Grade Perfeita (15 colunas)
                colunas_etiquetas = 15
                total_linhas_tabela = (len(df_filtrado) + colunas_etiquetas - 1) // colunas_etiquetas

                tabela = doc.add_table(rows=total_linhas_tabela, cols=colunas_etiquetas)
                tabela.autofit = False

                for col in tabela.columns:
                    col.width = Cm(1.9)
                for linha in tabela.rows:
                    linha.height = Cm(0.9)

                # 🟢 2. LOOP COM ÍNDICE RESETADO (Garante o alinhamento da grade)
                for index, row in df_filtrado.reset_index(drop=True).iterrows():
                    
                    r = index // colunas_etiquetas
                    c = index % colunas_etiquetas
                    celula = tabela.cell(r, c)
                    celula.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    pais = extrair_valor(row, map_pais).upper()
                    estado = extrair_valor(row, map_estado)
                    mun = extrair_valor(row, map_mun)
                    loc = extrair_valor(row, map_loc)
                    
                    lat = remover_milesimos_segundos(extrair_valor(row, map_lat))
                    lon = remover_milesimos_segundos(extrair_valor(row, map_lon))
                    coords = f"{lat} {lon}".strip()
                    
                    alt_str = extrair_valor(row, map_alt).replace("m", "").replace(".0", "").strip()
                    alt_formatada = f"{alt_str}m, " if alt_str else ""
                    
                    data_formatada = formatar_data_romana(extrair_valor(row, map_dt_ini), extrair_valor(row, map_dt_fim))
                    
                    isca = extrair_valor(row, map_isca)
                    isca_formatada = f"{isca}. " if isca else ""
                    
                    coletor = extrair_valor(row, map_coletor)
                    amostra = extrair_valor(row, map_id)
                    titulo_tax = extrair_valor(row, map_titulo_tax)

                    localizacao = f"{pais}: {estado}, {mun}, {loc}."
                    corpo_etiqueta = f"{localizacao} {coords}, {alt_formatada}{data_formatada}. {isca_formatada}{coletor}. {amostra}."
                    corpo_etiqueta = corpo_etiqueta.replace(" ,", ",").replace(" .", ".").replace("  ", " ").strip()

                    if exibir_titulo_tax and titulo_tax:
                        p_tit = celula.paragraphs[0]
                        p_tit.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        p_tit.paragraph_format.space_before, p_tit.paragraph_format.space_after = Pt(3), Pt(1)
                        p_tit.paragraph_format.line_spacing = 1.0
                        run_tit = p_tit.add_run(titulo_tax)
                        run_tit.italic = True
                        run_tit.font.name, run_tit.font.size = 'Arial', Pt(4)
                        p_corpo = celula.add_paragraph()
                    else:
                        p_corpo = celula.paragraphs[0]
                        p_corpo.paragraph_format.space_before = Pt(3)

                    p_corpo.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p_corpo.paragraph_format.line_spacing = 1.0
                    p_corpo.paragraph_format.space_after = Pt(3)
                    p_corpo.paragraph_format.left_indent = Cm(0)
                    p_corpo.paragraph_format.right_indent = Cm(0)

                    run_corpo = p_corpo.add_run(corpo_etiqueta)
                    run_corpo.font.name = 'Arial'
                    run_corpo.font.size = Pt(4)
                
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                
                st.success(f"🎉 Etiquetas geradas perfeitamente para {len(df_filtrado)} amostras!")
                st.download_button(
                    label="⬇️ Baixar Etiquetas em Word", 
                    data=buffer, 
                    file_name="Etiquetas_CEMT.docx", 
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
# ==========================================
    # 💬 ABA DE FEEDBACK (INTEGRADA AO GOOGLE SHEETS)
    # ==========================================
    with st.sidebar.expander("💬 Feedback e Sugestões", expanded=False):
        st.write("Encontrou algum problema ou tem sugestões de melhoria?")
        
        with st.form("form_feedback", clear_on_submit=True):
            nome_user = st.text_input("Seu nome (opcional):")
            email_user = st.text_input("Seu e-mail (opcional):")
            mensagem_user = st.text_area("Sua sugestão ou crítica:", placeholder="Digite aqui...")
            
            btn_enviar = st.form_submit_button("📩 Enviar Feedback")
            
            if btn_enviar:
                if not mensagem_user.strip():
                    st.warning("⚠️ Escreva uma mensagem antes de enviar.")
                else:
                    # 🟢 COLE O SEU URL DO APPS SCRIPT AQUI:
                    URL_WEB_APP = "https://script.google.com/macros/s/AKfycbz1nZslEYiH98x82DJo4YKM7NV20Gtfv3FhUlBQPGYWWK3ub0rrgZ1qr_4wrw4UBBg4/exec"
                    
                    payload = {
                        "nome": nome_user if nome_user else "Anônimo",
                        "email": email_user if email_user else "Não informado",
                        "mensagem": mensagem_user
                    }
                    
                    try:
                        # 🟢 Usamos json.dumps e headers explícitos para garantir a entrega
                        resposta = requests.post(
                            URL_WEB_APP, 
                            data=json.dumps(payload),
                            headers={"Content-Type": "application/json"}
                        )
                        
                        if resposta.status_code == 200:
                            st.success("✅ Feedback registrado na planilha com sucesso! Obrigado.")
                        else:
                            st.error(f"❌ Erro ao salvar (Código HTTP {resposta.status_code}).")
                    except Exception as e:
                        st.error("❌ Falha na conexão ao enviar o feedback.")